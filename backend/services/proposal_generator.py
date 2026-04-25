"""
יצירת הצעת מחיר בפורמט DOCX — 14 בלוקים קבועים עם slots.
RTL מלא, גופן David.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
from typing import Dict, List
import io
import os

from utils.template_blocks import (
    COMPANY_HEADER, COMPANY_INFO_FIXED, WORK_METHOD_FIXED,
    SCOPE_PLANNING_FIXED, SCOPE_EXECUTION_FIXED, SIGNATURE_FIXED,
    SPECIAL_CONDITIONS,
)
from utils.helpers import format_currency, format_currency_per_unit


def _set_rtl(paragraph):
    """הגדרת RTL לפסקה."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_rtl_paragraph(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    _set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "David"


def _add_section_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    _set_rtl(p)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "David"


def _add_table_rtl(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # כותרת
    hdr_row = table.rows[0]
    for i, header in enumerate(reversed(headers)):
        cell = hdr_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # שורות
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(reversed(row_data)):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_proposal_docx(
    project: object,
    calc_result: Dict,
    special_conditions: List[str],
) -> bytes:
    """
    יצירת קובץ DOCX מלא.
    מחזיר bytes לשמירה / הורדה.
    """
    doc = Document()

    # הגדרות RTL ברמת המסמך
    section = doc.sections[0]
    section.right_margin = Cm(2)
    section.left_margin = Cm(2)

    today = date.today().strftime("%d/%m/%Y")
    phase_costs = calc_result.get("phase_costs", [])
    schedule = calc_result.get("schedule", {})
    flags = calc_result.get("flags", {})
    blended_total = calc_result.get("blended_total", 0)
    blended_per_unit = calc_result.get("blended_per_unit", 0)
    price_range_low = calc_result.get("price_range_low", 0)
    price_range_high = calc_result.get("price_range_high", 0)
    ref_per_unit = calc_result.get("reference_price_per_unit", 0)
    total_months = schedule.get("total_months", 0) if isinstance(schedule, dict) else schedule.total_months

    # [B01] כותרת ופרטי מזמין
    _add_section_title(doc, "AM הנדסה")
    _add_rtl_paragraph(doc, "ניהול | תיאום תכנון | פיקוח")
    _add_rtl_paragraph(doc, "רח' המגשימים 20, פתח תקווה | טל: 03-5329221 | www.am-eng.net")
    doc.add_paragraph()
    _add_rtl_paragraph(doc, f"תאריך: {today}")
    _add_rtl_paragraph(doc, f"לכבוד: {project.client_name}")
    _add_rtl_paragraph(doc, "א.נ.")
    doc.add_paragraph()

    # [B02] הנדון
    _add_rtl_paragraph(
        doc,
        f"הנדון: הצעת מחיר לניהול, תיאום תכנון ופיקוח – {project.project_name}, {project.location_city}",
        bold=True,
        size=12,
    )
    doc.add_paragraph()

    # [B03] תיאור הפרויקט
    _add_section_title(doc, "תיאור הפרויקט")
    floors_underground = project.num_floors_underground or 0
    num_buildings = project.num_buildings or 1
    _add_rtl_paragraph(
        doc,
        f"הפרויקט כולל {project.num_units:,} יחידות דיור, {num_buildings} מבנה, "
        f"{project.num_floors_above} קומות עיליות"
        + (f" ו-{floors_underground:.0f} קומות מרתף" if floors_underground else "")
        + f". סוג הפרויקט: {project.project_type}.",
    )
    doc.add_paragraph()

    # [B04] פרטי חברת AM הנדסה
    _add_section_title(doc, "חברת AM הנדסה")
    _add_rtl_paragraph(doc, COMPANY_INFO_FIXED)
    doc.add_paragraph()

    # [B05] שיטת העבודה
    _add_section_title(doc, "שיטת העבודה")
    _add_rtl_paragraph(doc, WORK_METHOD_FIXED)
    doc.add_paragraph()

    # [B06] תכולת עבודה — תכנון
    _add_section_title(doc, "תכולת העבודה ותחומי אחריות — תכנון")
    _add_rtl_paragraph(doc, SCOPE_PLANNING_FIXED)
    doc.add_paragraph()

    # [B07] תכולת עבודה — ביצוע
    _add_section_title(doc, "תכולת העבודה ותחומי אחריות — ביצוע")
    _add_rtl_paragraph(doc, SCOPE_EXECUTION_FIXED)
    doc.add_paragraph()

    # [B08] לוח זמנים
    _add_section_title(doc, "לוחות זמנים משוערים")
    schedule_dict = schedule if isinstance(schedule, dict) else schedule.dict()
    phase_name_map = {
        "planning": "תכנון ורישוי",
        "excavation": "חפירה ודיפון",
        "underground": "שלד תת\"ק",
        "above_ground": "שלד עילי",
        "finishes": "גמרים / מעטפת / מערכות",
        "handover": "מסירות / טופס 4",
    }
    schedule_rows = []
    cumulative = 0
    for key, name in phase_name_map.items():
        months = schedule_dict.get(key, 0)
        if key in ("total_months", "source_note"):
            continue
        cumulative += months
        schedule_rows.append([name, str(months), str(cumulative)])

    _add_table_rtl(
        doc,
        ["פעילות", "משך (חודשים)", "מצטבר"],
        schedule_rows,
    )
    _add_rtl_paragraph(doc, f"סה\"כ משך הפרויקט: {total_months} חודשים")
    _add_rtl_paragraph(doc, "הערה: לוח הזמנים בהערכה בלבד.")
    doc.add_paragraph()

    # [B09] הצעת המחיר
    _add_section_title(doc, "הצעת המחיר")
    _add_rtl_paragraph(doc, "שכ\"ט יחולק על פי השלבים הבאים:")
    doc.add_paragraph()

    pricing_rows = []
    for phase in phase_costs:
        if isinstance(phase, dict):
            pricing_rows.append([
                phase["phase_name"],
                str(phase["months"]),
                format_currency(phase["adjusted_rate"]),
                format_currency(phase["phase_total"]),
            ])
        else:
            pricing_rows.append([
                phase.phase_name,
                str(phase.months),
                format_currency(phase.adjusted_rate),
                format_currency(phase.phase_total),
            ])

    _add_table_rtl(
        doc,
        ["שלב", "חודשים", "תעריף חודשי", "עלות שלב"],
        pricing_rows,
    )
    doc.add_paragraph()
    _add_rtl_paragraph(doc, f"הצעה כוללת מומלצת: {format_currency(blended_total)}", bold=True)
    _add_rtl_paragraph(doc, f"מחיר ליח\"ד: {format_currency_per_unit(blended_per_unit)}")
    _add_rtl_paragraph(doc, f"טווח מומלץ: {format_currency_per_unit(price_range_low)} – {format_currency_per_unit(price_range_high)}")
    if ref_per_unit:
        _add_rtl_paragraph(doc, f"מחיר ייחוס מפרויקטים דומים: {format_currency_per_unit(ref_per_unit)}")
    doc.add_paragraph()

    # [B10] תנאים מיוחדים
    if special_conditions:
        _add_section_title(doc, "הערות מיוחדות")
        for cond_key in special_conditions:
            cond_text = SPECIAL_CONDITIONS.get(cond_key, "")
            if cond_text:
                _add_rtl_paragraph(doc, f"• {cond_text}")
        doc.add_paragraph()

    # [B11] הערות כלליות
    _add_section_title(doc, "הערות")
    index_type = project.index_type or "תשומות בניה"
    vat_text = "כולל מע\"מ" if project.includes_vat else "אינו כולל מע\"מ"
    _add_rtl_paragraph(doc, f"• השכר {vat_text}.")
    _add_rtl_paragraph(doc, f"• שכ\"ט יהיה צמוד למדד {index_type} הידוע ביום הגשת ההצעה.")
    if project.notes_pricing:
        _add_rtl_paragraph(doc, f"• {project.notes_pricing}")
    doc.add_paragraph()

    # [B12] חתימה
    _add_rtl_paragraph(doc, SIGNATURE_FIXED)

    # שמירה ל-bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
